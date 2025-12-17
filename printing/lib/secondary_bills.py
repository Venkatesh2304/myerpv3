from docx import Document
from docx.shared import Pt, Cm
from prettytable import PrettyTable, ALL
from typing import Callable, Dict, Optional, List, Any
from jinja2 import Environment, FileSystemLoader
import weasyprint
import os
import base64

class SecondaryBillGenerator:
    def generate(self, txt_file_path: str, output_docx_path: str, barcode_generator: Callable, config: Optional[Dict]):
        document = Document()
        self._process_file(txt_file_path, document, barcode_generator, config)
        document.save(output_docx_path)

    def _process_file(self, file_path: str, document: Document, barcode_generator: Callable, config: Dict):
        with open(file_path, 'r') as f:
            content = f.read()
        
        lines = content.split('\n')
        values = ['Region', 'Invoice No', 'Invoice Date', 'Retailer PAN']
        first, last, billval = [], [], []
        invoice = []
        name = []

        for i in range(len(lines)):
            line = lines[i]
            if 'Invoice No ' in line and config['secname'] in line:
                first.append(i)
                invoice.append(line)
            if 'Time of Billing ' in line:
                last.append(i)
            if 'Bill Amount' in line:
                billval.append(line)
            if 'Retailer ' in line and 'Name' in line and config['secadd'] in line:
                name.append(line)


        self._setup_document_styles(document)


        for i in range(len(first)):

            y1 = lines[first[i]:last[i]+1]
            bill_text = ""
            for j in y1:
                bill_text += j + '\n'
                l = 0
                j1 = j
                for t in values:
                    if t in j:
                        l = 1
                        if 'Time' in j:
                            j1 = j
                        else:
                            j1 = j.split(t)[0]
                
                if l == 0:
                    document.add_paragraph(j)
                else:
                    document.add_paragraph(j1)

            billvalue = billval[i]
            billvalue1 = billvalue.split('Bill')[0]
            billvalue2 = 'Bill' + billvalue.split('Bill')[1]
            document.add_paragraph(billvalue1)
            
            paragraph1 = document.add_paragraph()
            
            # Safer parsing
            try:
                inv_part = invoice[i].split('Invoice')[1].split(':')[1]
                name_part = name[i].split(':')[1]
                amt_part = billvalue2.split(':')[1]
                imp = f"{inv_part}*{name_part}*Amt :{amt_part}"
                imp = ' '.join(imp.split())
                run = paragraph1.add_run('  ' + '   '.join(imp.split('*')))


            except IndexError:
                run = paragraph1.add_run("Error parsing invoice details")

            run.font.size = Pt(12)
            run.bold = True
            
            paragraph3 = document.add_paragraph().add_run(' ' * 10 + 'Signature')
            # paragraph3.alignment = 2 # WD_ALIGN_PARAGRAPH.RIGHT is 2
            # But run doesn't have alignment, paragraph does.
            document.paragraphs[-1].alignment = 2 # Right align
            paragraph3.font.size = Pt(12)
            paragraph3.bold = True

            # Barcode
            try:
                inum = invoice[i].split('Invoice')[1].split(':')[1].strip()
                barcode_img = barcode_generator(inum)
                barcode = document.add_picture(barcode_img)
                barcode.width = Cm(2.5)
                barcode.height = Cm(2.5)
            except Exception as e:
                print(f"Error adding barcode: {e}")

            lines_per_page = 15
            if i % 2 == 0:
                document.add_paragraph('\n' * lines_per_page)
            else:
                document.add_page_break()

    def _setup_document_styles(self, document: Document):
        style = document.styles['Normal']
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(9)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

        for section in document.sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(0.5)
            section.right_margin = Cm(0.5)

class SecondaryBillGeneratorWeasy:
    def generate(self, txt_file_path: str, output_pdf_path: str, barcode_generator: Callable, config: Optional[Dict], html_output_path: Optional[str] = None):
        data = self._process_file(txt_file_path, barcode_generator, config)
        self._render_pdf(data, output_pdf_path, config, html_output_path)

    def _process_file(self, file_path: str, barcode_generator: Callable, config: Dict) -> Dict[str, Any]:
        with open(file_path, 'r') as f:
            content = f.read()
        
        lines = content.split('\n')
        values = ['Region', 'Invoice No', 'Invoice Date', 'Retailer PAN']
        first, last, billval = [], [], []
        invoice = []
        name = []

        # Exact extraction logic from SecondaryBillGenerator
        for i in range(len(lines)):
            line = lines[i]
            # Note: config['secname'] and config['secadd'] are expected to be in the config
            # If not provided, we might need defaults or handle error. 
            # For now assuming they are provided as per existing code contract.
            secname = config.get('secname', '')
            secadd = config.get('secadd','')
            
            if 'Invoice No ' in line and secname in line:
                first.append(i)
                invoice.append(line)
            if 'Time of Billing ' in line:
                last.append(i)
            if 'Bill Amount' in line:
                billval.append(line)
            if 'Retailer ' in line and 'Name' in line and secadd in line:
                name.append(line)


        pages = []

        for i in range(len(first)):
            page_lines = []
            y1 = lines[first[i]:last[i]+1]
            
            for j in y1:
                # Logic to split lines based on 'values' keys
                # This logic in original code seems to be trying to handle cases where 
                # multiple fields are on the same line but maybe not separated cleanly?
                # Or maybe it's trying to wrap?
                # Original:
                # for t in values: if t in j: l=1...
                # if l==0: doc.add(j) else: doc.add(j1)
                
                # Let's replicate the logic to get the 'displayable' line
                l = 0
                j1 = j
                for t in values:
                    if t in j:
                        l = 1
                        if 'Time' in j:
                            j1 = j
                        else:
                            j1 = j.split(t)[0]
                
                if l == 0:
                    page_lines.append(j)
                else:
                    page_lines.append(j1)

            # Bill Value logic
            billvalue = billval[i]
            billvalue1 = billvalue.split('Bill')[0]
            billvalue2 = 'Bill' + billvalue.split('Bill')[1]
            page_lines.append(billvalue1)

            # Invoice Details Line (The bold line with Amt)
            invoice_details = ""
            try:
                inv_part = invoice[i].split('Invoice')[1].split(':')[1]
                inum = inv_part.strip()
                name_part = name[i].split(':')[1]
                amt_part = billvalue2.split(':')[1]
                imp = f"{inv_part}*{name_part}*Amt :{amt_part}"
                imp = ' '.join(imp.split())
                invoice_details = '  ' + '   '.join(imp.split('*'))


                # Generate barcode
                barcode_data = None
                if barcode_generator:
                    try:
                        img_stream = barcode_generator(inum)
                        if img_stream:
                            if hasattr(img_stream, 'getvalue'):
                                img_data = img_stream.getvalue()
                            elif hasattr(img_stream, 'read'):
                                img_stream.seek(0)
                                img_data = img_stream.read()
                            else:
                                # Assume bytes
                                img_data = img_stream
                            
                            b64_data = base64.b64encode(img_data).decode('utf-8')
                            # Assuming PNG for now, but could be JPEG. 
                            # The browser/weasyprint will likely handle it if we just say image/png or similar, 
                            # or we can try to detect. 
                            # For QR codes it's usually PNG.
                            barcode_data = f"data:image/png;base64,{b64_data}"
                    except Exception as e:
                        print(f"Error processing barcode for {inum}: {e}")
            except IndexError:
                invoice_details = "Error parsing invoice details"
                inum = ""
                barcode_data = None
            except Exception as e :
                print(e)
                
            pages.append({
                'lines': page_lines,
                'invoice_details': invoice_details,
                'inum': inum,
                'barcode_data': barcode_data,
                'is_last_page': True 
            })


        return {'pages': pages}

    def _render_pdf(self, data: Dict[str, Any], output_path: str, config: Optional[Dict], html_output_path: Optional[str] = None):
        env = Environment(loader=FileSystemLoader(os.path.join(os.path.dirname(__file__), '../templates')))
        template = env.get_template('secondary_bill.html')
        
        # We need to handle barcode generation here if possible, or pass a helper.
        # Since we don't have the barcode generator in the template context easily unless we pre-generate.
        # Let's assume we don't have barcode for now or we mock it.
        # The user passed barcode_generator to generate().
        
        # Wait, I need to pass barcode_svg to template.
        # I'll update data with barcode.
        
        # Default lines if not in config
        lines_spacing =  24

        html_out = template.render(pages=data['pages'], barcode_svg=None, lines_spacing=lines_spacing)

        
        if html_output_path:
            with open(html_output_path, 'w') as f:
                f.write(html_out)

        # WeasyPrint
        weasyprint.HTML(string=html_out).write_pdf(output_path)

